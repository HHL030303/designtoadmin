from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/wanhao/byy/designtoadmin")
OUTPUT_ROOT = ROOT / "outputs/软著材料"
SOURCE_ROOT = OUTPUT_ROOT / "backend_source_60_pages"
REFERENCE_DOC = Path(
    "/Users/wanhao/Library/Containers/com.tencent.WeWorkMac/Data/Documents/Profiles/"
    "32864579E46745C829576A15DA24F8F1/Caches/Files/2026-05/"
    "ef9e7adcd108fda8340021b210d747c5/武汉不一样教育科技有限公司10件证书+材料/"
    "武汉不一样教育科技有限公司10件证书+材料/5件软著-智慧物联网高端培训信息管控软件/"
    "3 数据化PPT演示设计应用管理平台/数据化PPT演示设计应用管理平台_源代码.docx"
)


def code_block(text: str) -> str:
    return dedent(text).strip("\n") + "\n"


EXISTING_FILES = sorted((OUTPUT_ROOT / "backend_source_30_pages").rglob("*.py"))


NEW_MODULES: dict[str, str] = {
    "app/models/project.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy import Boolean, Datetime, ForeignKey, Integer, String, Text
        from sqlalchemy.orm import Mapped, mapped_column, relationship

        from app.core.database import Base


        class ProjectRecord(Base):
            __tablename__ = "project_record"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            project_code: Mapped[str] = mapped_column(String(32), unique=True, index=True)
            project_name: Mapped[str] = mapped_column(String(100), index=True)
            description: Mapped[str | None] = mapped_column(Text(), nullable=True)
            status: Mapped[str] = mapped_column(String(20), default="enabled")
            created_at: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)
            updated_at: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)

            workflow_templates: Mapped[list["WorkflowTemplateRecord"]] = relationship(
                back_populates="project",
                cascade="all, delete-orphan",
            )
            field_configs: Mapped[list["TaskFieldConfig"]] = relationship(
                back_populates="project",
                cascade="all, delete-orphan",
            )
            members: Mapped[list["ProjectMemberRecord"]] = relationship(
                back_populates="project",
                cascade="all, delete-orphan",
            )


        class WorkflowTemplateRecord(Base):
            __tablename__ = "workflow_template_record"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            project_id: Mapped[int] = mapped_column(ForeignKey("project_record.id", ondelete="CASCADE"))
            template_name: Mapped[str] = mapped_column(String(100))
            order_type: Mapped[str] = mapped_column(String(20), default="new")
            is_default: Mapped[bool] = mapped_column(Boolean(), default=False)
            status: Mapped[str] = mapped_column(String(20), default="enabled")
            created_at: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)
            updated_at: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)

            project: Mapped[ProjectRecord] = relationship(back_populates="workflow_templates")
            stages: Mapped[list["WorkflowStageRecord"]] = relationship(
                back_populates="template",
                cascade="all, delete-orphan",
            )


        class WorkflowStageRecord(Base):
            __tablename__ = "workflow_stage_record"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            template_id: Mapped[int] = mapped_column(
                ForeignKey("workflow_template_record.id", ondelete="CASCADE")
            )
            stage_name: Mapped[str] = mapped_column(String(100))
            role_code: Mapped[str] = mapped_column(String(50))
            sort_value: Mapped[int] = mapped_column(Integer(), default=1)
            can_assign: Mapped[bool] = mapped_column(Boolean(), default=False)
            can_skip: Mapped[bool] = mapped_column(Boolean(), default=False)
            requires_file_upload: Mapped[bool] = mapped_column(Boolean(), default=False)
            requires_validation: Mapped[bool] = mapped_column(Boolean(), default=False)
            triggers_package: Mapped[bool] = mapped_column(Boolean(), default=False)

            template: Mapped[WorkflowTemplateRecord] = relationship(back_populates="stages")


        class TaskFieldConfig(Base):
            __tablename__ = "task_field_config"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            project_id: Mapped[int] = mapped_column(ForeignKey("project_record.id", ondelete="CASCADE"))
            field_key: Mapped[str] = mapped_column(String(50), index=True)
            field_name: Mapped[str] = mapped_column(String(50))
            field_type: Mapped[str] = mapped_column(String(20))
            required: Mapped[bool] = mapped_column(Boolean(), default=False)
            searchable: Mapped[bool] = mapped_column(Boolean(), default=False)
            sort_value: Mapped[int] = mapped_column(Integer(), default=1)
            status: Mapped[str] = mapped_column(String(20), default="enabled")

            project: Mapped[ProjectRecord] = relationship(back_populates="field_configs")


        class ProjectMemberRecord(Base):
            __tablename__ = "project_member_record"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            project_id: Mapped[int] = mapped_column(ForeignKey("project_record.id", ondelete="CASCADE"))
            user_id: Mapped[int] = mapped_column(ForeignKey("user_account.id", ondelete="CASCADE"))
            role_code: Mapped[str] = mapped_column(String(50))
            role_name: Mapped[str] = mapped_column(String(50))
            created_at: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)

            project: Mapped[ProjectRecord] = relationship(back_populates="members")
        """
    ),
    "app/models/statistics.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy import Datetime, ForeignKey, Integer, String
        from sqlalchemy.orm import Mapped, mapped_column

        from app.core.database import Base


        class TaskDailySnapshot(Base):
            __tablename__ = "task_daily_snapshot"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            task_code: Mapped[str] = mapped_column(String(32), index=True)
            project_code: Mapped[str] = mapped_column(String(32), index=True)
            subject: Mapped[str] = mapped_column(String(50), index=True)
            grade: Mapped[str] = mapped_column(String(20), index=True)
            current_stage: Mapped[str] = mapped_column(String(50), index=True)
            current_owner: Mapped[str] = mapped_column(String(100))
            overdue_count: Mapped[int] = mapped_column(Integer(), default=0)
            pending_count: Mapped[int] = mapped_column(Integer(), default=0)
            processing_count: Mapped[int] = mapped_column(Integer(), default=0)
            completed_count: Mapped[int] = mapped_column(Integer(), default=0)
            snapshot_date: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)


        class DesignerWorkloadSnapshot(Base):
            __tablename__ = "designer_workload_snapshot"

            id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
            user_id: Mapped[int] = mapped_column(ForeignKey("user_account.id", ondelete="CASCADE"))
            user_name: Mapped[str] = mapped_column(String(50), index=True)
            role_code: Mapped[str] = mapped_column(String(50))
            task_total: Mapped[int] = mapped_column(Integer(), default=0)
            assigned_pages: Mapped[int] = mapped_column(Integer(), default=0)
            completed_pages: Mapped[int] = mapped_column(Integer(), default=0)
            snapshot_date: Mapped[datetime] = mapped_column(Datetime(), default=datetime.utcnow)
        """
    ),
    "app/schemas/project.py": code_block(
        """
        from pydantic import BaseModel, Field


        class ProjectCreateRequest(BaseModel):
            project_code: str = Field(min_length=2, max_length=32)
            project_name: str = Field(min_length=2, max_length=100)
            description: str | None = None
            status: str = "enabled"


        class ProjectUpdateRequest(BaseModel):
            project_name: str = Field(min_length=2, max_length=100)
            description: str | None = None
            status: str = "enabled"


        class WorkflowStagePayload(BaseModel):
            stage_name: str
            role_code: str
            sort_value: int
            can_assign: bool = False
            can_skip: bool = False
            requires_file_upload: bool = False
            requires_validation: bool = False
            triggers_package: bool = False


        class WorkflowTemplatePayload(BaseModel):
            template_name: str
            order_type: str = "new"
            is_default: bool = False
            status: str = "enabled"
            stages: list[WorkflowStagePayload]


        class FieldConfigPayload(BaseModel):
            field_key: str
            field_name: str
            field_type: str
            required: bool = False
            searchable: bool = False
            sort_value: int = 1
            status: str = "enabled"


        class ProjectMemberPayload(BaseModel):
            user_id: int
            role_code: str
            role_name: str
        """
    ),
    "app/schemas/admin.py": code_block(
        """
        from pydantic import BaseModel, EmailStr, Field


        class AdminUserCreateRequest(BaseModel):
            email: EmailStr
            password: str = Field(min_length=6, max_length=50)
            name: str = Field(min_length=2, max_length=50)
            role_code: str = "planner"
            enabled: bool = True


        class AdminUserUpdateRequest(BaseModel):
            name: str = Field(min_length=2, max_length=50)
            role_code: str = "planner"
            enabled: bool = True


        class RoleConfigPayload(BaseModel):
            code: str
            name: str
            description: str
            scope: str
            view_access: list[str]
        """
    ),
    "app/schemas/file_storage.py": code_block(
        """
        from datetime import datetime

        from pydantic import BaseModel


        class UploadPolicyResponse(BaseModel):
            bucket: str
            region: str
            secret_id: str
            token: str
            expired_at: datetime


        class UploadedFilePayload(BaseModel):
            task_code: str
            file_name: str
            file_ext: str | None = None
            storage_key: str | None = None
            checksum: str | None = None
            category: str = "general"
            uploaded_by: str


        class PackageTriggerRequest(BaseModel):
            task_code: str
            operator_name: str
            archive_name: str
        """
    ),
    "app/repositories/project_repository.py": code_block(
        """
        from collections.abc import Sequence

        from sqlalchemy import Select, func, select
        from sqlalchemy.orm import Session, selectinload

        from app.models.project import ProjectRecord


        class ProjectRepository:
            def __init__(self, db: Session) -> None:
                self.db = db

            def list(self) -> Sequence[ProjectRecord]:
                stmt: Select[tuple[ProjectRecord]] = (
                    select(ProjectRecord)
                    .options(
                        selectinload(ProjectRecord.workflow_templates),
                        selectinload(ProjectRecord.field_configs),
                        selectinload(ProjectRecord.members),
                    )
                    .order_by(ProjectRecord.id.desc())
                )
                return self.db.execute(stmt).scalars().all()

            def get_by_code(self, project_code: str) -> ProjectRecord | None:
                stmt = (
                    select(ProjectRecord)
                    .options(
                        selectinload(ProjectRecord.workflow_templates),
                        selectinload(ProjectRecord.field_configs),
                        selectinload(ProjectRecord.members),
                    )
                    .where(ProjectRecord.project_code == project_code)
                )
                return self.db.execute(stmt).scalar_one_or_none()

            def get_total(self) -> int:
                stmt = select(func.count(ProjectRecord.id))
                total = self.db.execute(stmt).scalar_one_or_none()
                return total or 0

            def add(self, record: ProjectRecord) -> ProjectRecord:
                self.db.add(record)
                self.db.flush()
                self.db.refresh(record)
                return record

            def save(self) -> None:
                self.db.commit()
        """
    ),
    "app/repositories/statistics_repository.py": code_block(
        """
        from collections.abc import Sequence
        from datetime import datetime

        from sqlalchemy import Select, func, select
        from sqlalchemy.orm import Session

        from app.models.statistics import DesignerWorkloadSnapshot, TaskDailySnapshot


        class StatisticsRepository:
            def __init__(self, db: Session) -> None:
                self.db = db

            def list_task_snapshots(
                self,
                date_from: datetime | None = None,
                date_to: datetime | None = None,
            ) -> Sequence[TaskDailySnapshot]:
                stmt: Select[tuple[TaskDailySnapshot]] = select(TaskDailySnapshot).order_by(
                    TaskDailySnapshot.snapshot_date.desc()
                )
                if date_from is not None:
                    stmt = stmt.where(TaskDailySnapshot.snapshot_date >= date_from)
                if date_to is not None:
                    stmt = stmt.where(TaskDailySnapshot.snapshot_date <= date_to)
                return self.db.execute(stmt).scalars().all()

            def list_designer_snapshots(self) -> Sequence[DesignerWorkloadSnapshot]:
                stmt = select(DesignerWorkloadSnapshot).order_by(
                    DesignerWorkloadSnapshot.snapshot_date.desc()
                )
                return self.db.execute(stmt).scalars().all()

            def add_task_snapshot(self, snapshot: TaskDailySnapshot) -> TaskDailySnapshot:
                self.db.add(snapshot)
                self.db.flush()
                self.db.refresh(snapshot)
                return snapshot

            def add_designer_snapshot(
                self,
                snapshot: DesignerWorkloadSnapshot,
            ) -> DesignerWorkloadSnapshot:
                self.db.add(snapshot)
                self.db.flush()
                self.db.refresh(snapshot)
                return snapshot

            def get_stage_count(self, stage_name: str) -> int:
                stmt = select(func.count(TaskDailySnapshot.id)).where(
                    TaskDailySnapshot.current_stage == stage_name
                )
                total = self.db.execute(stmt).scalar_one_or_none()
                return total or 0

            def save(self) -> None:
                self.db.commit()
        """
    ),
    "app/repositories/attachment_repository.py": code_block(
        """
        from collections.abc import Sequence

        from sqlalchemy import select
        from sqlalchemy.orm import Session

        from app.models.course import CourseAttachment


        class AttachmentRepository:
            def __init__(self, db: Session) -> None:
                self.db = db

            def list_by_task_code(self, task_code: str) -> Sequence[CourseAttachment]:
                stmt = (
                    select(CourseAttachment)
                    .join(CourseAttachment.course)
                    .where(CourseAttachment.course.has(task_code=task_code))
                    .order_by(CourseAttachment.id.desc())
                )
                return self.db.execute(stmt).scalars().all()

            def add(self, attachment: CourseAttachment) -> CourseAttachment:
                self.db.add(attachment)
                self.db.flush()
                self.db.refresh(attachment)
                return attachment

            def save(self) -> None:
                self.db.commit()
        """
    ),
    "app/services/project_service.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy.orm import Session

        from app.models.project import (
            ProjectMemberRecord,
            ProjectRecord,
            TaskFieldConfig,
            WorkflowStageRecord,
            WorkflowTemplateRecord,
        )
        from app.repositories.project_repository import ProjectRepository
        from app.schemas.project import (
            FieldConfigPayload,
            ProjectCreateRequest,
            ProjectMemberPayload,
            ProjectUpdateRequest,
            WorkflowTemplatePayload,
        )


        class ProjectService:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.repository = ProjectRepository(db)

            def list_projects(self) -> list[ProjectRecord]:
                return list(self.repository.list())

            def create_project(self, payload: ProjectCreateRequest) -> ProjectRecord:
                record = ProjectRecord(
                    project_code=payload.project_code,
                    project_name=payload.project_name,
                    description=payload.description,
                    status=payload.status,
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow(),
                )
                self.repository.add(record)
                self.repository.save()
                return record

            def update_project(
                self,
                project_code: str,
                payload: ProjectUpdateRequest,
            ) -> ProjectRecord:
                record = self.get_project(project_code)
                record.project_name = payload.project_name
                record.description = payload.description
                record.status = payload.status
                record.updated_at = datetime.utcnow()
                self.repository.save()
                return record

            def get_project(self, project_code: str) -> ProjectRecord:
                record = self.repository.get_by_code(project_code)
                if record is None:
                    raise ValueError("project_not_found")
                return record

            def replace_workflow_templates(
                self,
                project_code: str,
                payloads: list[WorkflowTemplatePayload],
            ) -> ProjectRecord:
                record = self.get_project(project_code)
                record.workflow_templates.clear()
                for payload in payloads:
                    template = WorkflowTemplateRecord(
                        template_name=payload.template_name,
                        order_type=payload.order_type,
                        is_default=payload.is_default,
                        status=payload.status,
                        created_at=datetime.utcnow(),
                        updated_at=datetime.utcnow(),
                    )
                    for stage_payload in payload.stages:
                        template.stages.append(
                            WorkflowStageRecord(
                                stage_name=stage_payload.stage_name,
                                role_code=stage_payload.role_code,
                                sort_value=stage_payload.sort_value,
                                can_assign=stage_payload.can_assign,
                                can_skip=stage_payload.can_skip,
                                requires_file_upload=stage_payload.requires_file_upload,
                                requires_validation=stage_payload.requires_validation,
                                triggers_package=stage_payload.triggers_package,
                            )
                        )
                    record.workflow_templates.append(template)
                record.updated_at = datetime.utcnow()
                self.repository.save()
                return record

            def replace_field_configs(
                self,
                project_code: str,
                payloads: list[FieldConfigPayload],
            ) -> ProjectRecord:
                record = self.get_project(project_code)
                record.field_configs.clear()
                for payload in payloads:
                    record.field_configs.append(
                        TaskFieldConfig(
                            field_key=payload.field_key,
                            field_name=payload.field_name,
                            field_type=payload.field_type,
                            required=payload.required,
                            searchable=payload.searchable,
                            sort_value=payload.sort_value,
                            status=payload.status,
                        )
                    )
                record.updated_at = datetime.utcnow()
                self.repository.save()
                return record

            def replace_project_members(
                self,
                project_code: str,
                payloads: list[ProjectMemberPayload],
            ) -> ProjectRecord:
                record = self.get_project(project_code)
                record.members.clear()
                for payload in payloads:
                    record.members.append(
                        ProjectMemberRecord(
                            user_id=payload.user_id,
                            role_code=payload.role_code,
                            role_name=payload.role_name,
                            created_at=datetime.utcnow(),
                        )
                    )
                record.updated_at = datetime.utcnow()
                self.repository.save()
                return record
        """
    ),
    "app/services/admin_service.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy.orm import Session

        from app.core.security import get_password_hash
        from app.models.user import UserAccount
        from app.repositories.user_repository import UserRepository
        from app.schemas.admin import AdminUserCreateRequest, AdminUserUpdateRequest


        class AdminService:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.repository = UserRepository(db)

            def list_users(self) -> list[UserAccount]:
                return list(self.repository.list_all())

            def create_user(self, payload: AdminUserCreateRequest) -> UserAccount:
                exists = self.repository.get_by_email(payload.email)
                if exists is not None:
                    raise ValueError("email_exists")
                user = UserAccount(
                    email=payload.email,
                    password_hash=get_password_hash(payload.password),
                    name=payload.name,
                    role_code=payload.role_code,
                    enabled=payload.enabled,
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow(),
                )
                self.repository.add(user)
                self.repository.save()
                return user

            def update_user(self, user_id: int, payload: AdminUserUpdateRequest) -> UserAccount:
                user = self.repository.get_by_id(user_id)
                if user is None:
                    raise ValueError("user_not_found")
                user.name = payload.name
                user.role_code = payload.role_code
                user.enabled = payload.enabled
                user.updated_at = datetime.utcnow()
                self.repository.save()
                return user

            def disable_user(self, user_id: int) -> UserAccount:
                user = self.repository.get_by_id(user_id)
                if user is None:
                    raise ValueError("user_not_found")
                user.enabled = False
                user.updated_at = datetime.utcnow()
                self.repository.save()
                return user
        """
    ),
    "app/services/file_storage_service.py": code_block(
        """
        from datetime import datetime, timedelta
        from uuid import uuid4

        from sqlalchemy.orm import Session

        from app.models.course import CourseAttachment, CourseLog
        from app.repositories.course_repository import CourseRepository
        from app.schemas.file_storage import PackageTriggerRequest, UploadPolicyResponse, UploadedFilePayload


        class FileStorageService:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.course_repository = CourseRepository(db)

            def get_upload_policy(self) -> UploadPolicyResponse:
                now = datetime.utcnow()
                return UploadPolicyResponse(
                    bucket="design-delivery-bucket",
                    region="ap-shanghai",
                    secret_id="TMP_SECRET_ID",
                    token="TMP_TOKEN",
                    expired_at=now + timedelta(minutes=30),
                )

            def record_upload(self, payload: UploadedFilePayload) -> dict[str, str]:
                task = self.course_repository.get_by_task_code(payload.task_code)
                if task is None:
                    raise ValueError("course_not_found")
                attachment = CourseAttachment(
                    course=task,
                    file_name=payload.file_name,
                    file_ext=payload.file_ext,
                    storage_key=payload.storage_key or f"{payload.task_code}/{uuid4().hex}",
                    category=payload.category,
                    uploaded_by=payload.uploaded_by,
                    uploaded_at=datetime.utcnow(),
                )
                task.attachments.append(attachment)
                task.logs.append(
                    CourseLog(
                        actor_name=payload.uploaded_by,
                        action_name="记录上传文件",
                        detail_text=f"已记录文件 {payload.file_name}",
                    )
                )
                self.course_repository.save()
                return {
                    "task_code": task.task_code,
                    "file_name": attachment.file_name,
                    "storage_key": attachment.storage_key or "",
                }

            def trigger_package(self, payload: PackageTriggerRequest) -> dict[str, str]:
                task = self.course_repository.get_by_task_code(payload.task_code)
                if task is None:
                    raise ValueError("course_not_found")
                task.logs.append(
                    CourseLog(
                        actor_name=payload.operator_name,
                        action_name="触发打包",
                        detail_text=f"已触发归档打包 {payload.archive_name}",
                    )
                )
                self.course_repository.save()
                return {
                    "task_code": task.task_code,
                    "archive_name": payload.archive_name,
                    "status": "packing",
                }
        """
    ),
    "app/services/statistics_service.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy.orm import Session

        from app.models.statistics import DesignerWorkloadSnapshot, TaskDailySnapshot
        from app.repositories.course_repository import CourseRepository
        from app.repositories.statistics_repository import StatisticsRepository


        class StatisticsService:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.course_repository = CourseRepository(db)
                self.statistics_repository = StatisticsRepository(db)

            def rebuild_task_snapshots(self, project_code: str) -> int:
                tasks, _total = self.course_repository.list(limit=5000)
                count = 0
                for task in tasks:
                    self.statistics_repository.add_task_snapshot(
                        TaskDailySnapshot(
                            task_code=task.task_code,
                            project_code=project_code,
                            subject=task.subject,
                            grade=task.grade,
                            current_stage=task.status,
                            current_owner=task.current_owner,
                            overdue_count=1 if task.overdue else 0,
                            pending_count=1 if task.status != "archived" else 0,
                            processing_count=1 if task.status not in {"archived", "packing"} else 0,
                            completed_count=1 if task.status == "archived" else 0,
                            snapshot_date=datetime.utcnow(),
                        )
                    )
                    count += 1
                self.statistics_repository.save()
                return count

            def rebuild_designer_snapshots(self) -> int:
                tasks, _total = self.course_repository.list(limit=5000)
                grouped: dict[str, DesignerWorkloadSnapshot] = {}
                for task in tasks:
                    owner = task.current_owner
                    if "设计师" not in owner:
                        continue
                    if owner not in grouped:
                        grouped[owner] = DesignerWorkloadSnapshot(
                            user_id=0,
                            user_name=owner,
                            role_code="designer",
                            task_total=0,
                            assigned_pages=0,
                            completed_pages=0,
                            snapshot_date=datetime.utcnow(),
                        )
                    grouped[owner].task_total += 1
                    grouped[owner].assigned_pages += task.total_page_count
                    if task.status == "archived":
                        grouped[owner].completed_pages += task.total_page_count
                for snapshot in grouped.values():
                    self.statistics_repository.add_designer_snapshot(snapshot)
                self.statistics_repository.save()
                return len(grouped)

            def get_project_overview(self, project_code: str) -> dict[str, int]:
                snapshots = self.statistics_repository.list_task_snapshots()
                filtered = [item for item in snapshots if item.project_code == project_code]
                return {
                    "task_total": len(filtered),
                    "processing_total": sum(item.processing_count for item in filtered),
                    "completed_total": sum(item.completed_count for item in filtered),
                    "overdue_total": sum(item.overdue_count for item in filtered),
                }
        """
    ),
    "app/services/validation_service.py": code_block(
        """
        from app.schemas.project import FieldConfigPayload, WorkflowTemplatePayload


        class ValidationService:
            def validate_field_configs(self, payloads: list[FieldConfigPayload]) -> None:
                if not payloads:
                    raise ValueError("field_configs_required")
                seen_keys: set[str] = set()
                for payload in payloads:
                    if payload.field_key in seen_keys:
                        raise ValueError("field_key_duplicated")
                    seen_keys.add(payload.field_key)
                    if payload.sort_value <= 0:
                        raise ValueError("field_sort_invalid")

            def validate_workflow_templates(
                self,
                payloads: list[WorkflowTemplatePayload],
            ) -> None:
                if not payloads:
                    raise ValueError("workflow_templates_required")
                default_count = sum(1 for item in payloads if item.is_default)
                if default_count > 1:
                    raise ValueError("workflow_default_duplicated")
                for payload in payloads:
                    if not payload.stages:
                        raise ValueError("workflow_stage_required")
                    sort_values = [stage.sort_value for stage in payload.stages]
                    if len(sort_values) != len(set(sort_values)):
                        raise ValueError("workflow_stage_sort_duplicated")

            def validate_page_assignment(
                self,
                total_page_count: int,
                primary_pages: int,
                assist_pages: list[int],
            ) -> None:
                assigned_total = primary_pages + sum(assist_pages)
                if assigned_total != total_page_count:
                    raise ValueError("page_assignment_mismatch")

            def validate_service_request(self, description: str, responsibility: str) -> None:
                if len(description.strip()) < 2:
                    raise ValueError("service_description_too_short")
                if responsibility not in {"设计责任", "内容责任", "客户需求变更", "其他"}:
                    raise ValueError("service_responsibility_invalid")
        """
    ),
    "app/services/audit_service.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy.orm import Session

        from app.models.course import CourseLog
        from app.repositories.course_repository import CourseRepository


        class AuditService:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.course_repository = CourseRepository(db)

            def append_task_log(
                self,
                task_code: str,
                actor_name: str,
                action_name: str,
                detail_text: str,
            ) -> dict[str, str]:
                task = self.course_repository.get_by_task_code(task_code)
                if task is None:
                    raise ValueError("course_not_found")
                task.logs.append(
                    CourseLog(
                        actor_name=actor_name,
                        action_name=action_name,
                        detail_text=detail_text,
                        created_at=datetime.utcnow(),
                    )
                )
                self.course_repository.save()
                return {
                    "task_code": task.task_code,
                    "action_name": action_name,
                    "actor_name": actor_name,
                }

            def rebuild_archive_logs(self, task_code: str) -> int:
                task = self.course_repository.get_by_task_code(task_code)
                if task is None:
                    raise ValueError("course_not_found")
                count = 0
                for attachment in task.attachments:
                    if attachment.category != "archive":
                        continue
                    task.logs.append(
                        CourseLog(
                            actor_name="system",
                            action_name="补录归档记录",
                            detail_text=f"发现归档文件 {attachment.file_name}",
                            created_at=datetime.utcnow(),
                        )
                    )
                    count += 1
                self.course_repository.save()
                return count
        """
    ),
    "app/services/report_service.py": code_block(
        """
        from datetime import datetime

        from sqlalchemy.orm import Session

        from app.repositories.course_repository import CourseRepository


        class ReportService:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.course_repository = CourseRepository(db)

            def build_task_export_rows(self) -> list[dict[str, object]]:
                tasks, _total = self.course_repository.list(limit=5000)
                rows: list[dict[str, object]] = []
                for task in tasks:
                    rows.append(
                        {
                            "任务编号": task.task_code,
                            "课件名称": task.title,
                            "学科": task.subject,
                            "年级": task.grade,
                            "当前阶段": task.status,
                            "当前负责人": task.current_owner,
                            "版本号": task.version_no,
                            "是否逾期": "是" if task.overdue else "否",
                            "归档时间": task.archived_at.isoformat() if task.archived_at else "",
                        }
                    )
                return rows

            def build_stage_summary(self) -> dict[str, int]:
                tasks, _total = self.course_repository.list(limit=5000)
                summary: dict[str, int] = {}
                for task in tasks:
                    summary[task.status] = summary.get(task.status, 0) + 1
                return summary

            def export_timestamp(self) -> str:
                return datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        """
    ),
    "app/api/routes/projects.py": code_block(
        """
        from fastapi import APIRouter, Depends, HTTPException

        from app.api.deps import DbSession, get_current_user
        from app.schemas.common import ApiResponse
        from app.schemas.project import (
            FieldConfigPayload,
            ProjectCreateRequest,
            ProjectMemberPayload,
            ProjectUpdateRequest,
            WorkflowTemplatePayload,
        )
        from app.services.project_service import ProjectService
        from app.services.validation_service import ValidationService


        router = APIRouter()


        @router.get("", response_model=ApiResponse[list[dict[str, object]]])
        def list_projects(
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[list[dict[str, object]]]:
            records = ProjectService(db).list_projects()
            return ApiResponse(
                data=[
                    {
                        "project_code": record.project_code,
                        "project_name": record.project_name,
                        "description": record.description,
                        "status": record.status,
                        "member_count": len(record.members),
                    }
                    for record in records
                ]
            )


        @router.post("", response_model=ApiResponse[dict[str, str]])
        def create_project(
            payload: ProjectCreateRequest,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            record = ProjectService(db).create_project(payload)
            return ApiResponse(data={"project_code": record.project_code, "status": record.status})


        @router.put("/{project_code}", response_model=ApiResponse[dict[str, str]])
        def update_project(
            project_code: str,
            payload: ProjectUpdateRequest,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            record = ProjectService(db).update_project(project_code, payload)
            return ApiResponse(data={"project_code": record.project_code, "status": record.status})


        @router.put("/{project_code}/workflow", response_model=ApiResponse[dict[str, str]])
        def replace_workflow(
            project_code: str,
            payloads: list[WorkflowTemplatePayload],
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            ValidationService().validate_workflow_templates(payloads)
            record = ProjectService(db).replace_workflow_templates(project_code, payloads)
            return ApiResponse(data={"project_code": record.project_code, "message": "workflow_saved"})


        @router.put("/{project_code}/fields", response_model=ApiResponse[dict[str, str]])
        def replace_fields(
            project_code: str,
            payloads: list[FieldConfigPayload],
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            ValidationService().validate_field_configs(payloads)
            record = ProjectService(db).replace_field_configs(project_code, payloads)
            return ApiResponse(data={"project_code": record.project_code, "message": "fields_saved"})


        @router.put("/{project_code}/members", response_model=ApiResponse[dict[str, str]])
        def replace_members(
            project_code: str,
            payloads: list[ProjectMemberPayload],
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            record = ProjectService(db).replace_project_members(project_code, payloads)
            return ApiResponse(data={"project_code": record.project_code, "message": "members_saved"})
        """
    ),
    "app/api/routes/admin.py": code_block(
        """
        from fastapi import APIRouter, Depends, HTTPException

        from app.api.deps import DbSession, get_current_user
        from app.schemas.admin import AdminUserCreateRequest, AdminUserUpdateRequest
        from app.schemas.common import ApiResponse
        from app.services.admin_service import AdminService


        router = APIRouter()


        @router.get("/users", response_model=ApiResponse[list[dict[str, object]]])
        def list_users(
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[list[dict[str, object]]]:
            users = AdminService(db).list_users()
            return ApiResponse(
                data=[
                    {
                        "id": user.id,
                        "email": user.email,
                        "name": user.name,
                        "role_code": user.role_code,
                        "enabled": user.enabled,
                    }
                    for user in users
                ]
            )


        @router.post("/users", response_model=ApiResponse[dict[str, object]])
        def create_user(
            payload: AdminUserCreateRequest,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, object]]:
            user = AdminService(db).create_user(payload)
            return ApiResponse(data={"id": user.id, "email": user.email, "role_code": user.role_code})


        @router.put("/users/{user_id}", response_model=ApiResponse[dict[str, object]])
        def update_user(
            user_id: int,
            payload: AdminUserUpdateRequest,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, object]]:
            user = AdminService(db).update_user(user_id, payload)
            return ApiResponse(data={"id": user.id, "name": user.name, "enabled": user.enabled})


        @router.post("/users/{user_id}/disable", response_model=ApiResponse[dict[str, object]])
        def disable_user(
            user_id: int,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, object]]:
            try:
                user = AdminService(db).disable_user(user_id)
            except ValueError as exc:
                raise HTTPException(status_code=404, detail="未找到用户") from exc
            return ApiResponse(data={"id": user.id, "enabled": user.enabled})
        """
    ),
    "app/api/routes/files.py": code_block(
        """
        from fastapi import APIRouter, Depends, HTTPException

        from app.api.deps import DbSession, get_current_user
        from app.schemas.common import ApiResponse
        from app.schemas.file_storage import PackageTriggerRequest, UploadedFilePayload
        from app.services.file_storage_service import FileStorageService


        router = APIRouter()


        @router.get("/upload-policy", response_model=ApiResponse[dict[str, object]])
        def get_upload_policy(
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, object]]:
            policy = FileStorageService(db).get_upload_policy()
            return ApiResponse(data=policy.model_dump())


        @router.post("/record", response_model=ApiResponse[dict[str, str]])
        def record_upload(
            payload: UploadedFilePayload,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            try:
                data = FileStorageService(db).record_upload(payload)
            except ValueError as exc:
                raise HTTPException(status_code=404, detail="未找到任务") from exc
            return ApiResponse(data=data)


        @router.post("/package", response_model=ApiResponse[dict[str, str]])
        def trigger_package(
            payload: PackageTriggerRequest,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, str]]:
            try:
                data = FileStorageService(db).trigger_package(payload)
            except ValueError as exc:
                raise HTTPException(status_code=404, detail="未找到任务") from exc
            return ApiResponse(data=data)
        """
    ),
    "app/api/routes/reports.py": code_block(
        """
        from fastapi import APIRouter, Depends

        from app.api.deps import DbSession, get_current_user
        from app.schemas.common import ApiResponse
        from app.services.report_service import ReportService


        router = APIRouter()


        @router.get("/tasks", response_model=ApiResponse[dict[str, object]])
        def export_tasks(
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, object]]:
            service = ReportService(db)
            return ApiResponse(
                data={
                    "generated_at": service.export_timestamp(),
                    "rows": service.build_task_export_rows(),
                }
            )


        @router.get("/stages", response_model=ApiResponse[dict[str, object]])
        def export_stage_summary(
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, object]]:
            service = ReportService(db)
            return ApiResponse(
                data={
                    "generated_at": service.export_timestamp(),
                    "summary": service.build_stage_summary(),
                }
            )
        """
    ),
    "app/api/routes/statistics.py": code_block(
        """
        from fastapi import APIRouter, Depends

        from app.api.deps import DbSession, get_current_user
        from app.schemas.common import ApiResponse
        from app.services.statistics_service import StatisticsService


        router = APIRouter()


        @router.post("/rebuild/task-snapshots", response_model=ApiResponse[dict[str, int]])
        def rebuild_task_snapshots(
            project_code: str,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, int]]:
            count = StatisticsService(db).rebuild_task_snapshots(project_code)
            return ApiResponse(data={"rebuilt_count": count})


        @router.post("/rebuild/designer-snapshots", response_model=ApiResponse[dict[str, int]])
        def rebuild_designer_snapshots(
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, int]]:
            count = StatisticsService(db).rebuild_designer_snapshots()
            return ApiResponse(data={"rebuilt_count": count})


        @router.get("/overview", response_model=ApiResponse[dict[str, int]])
        def get_project_overview(
            project_code: str,
            db: DbSession = Depends(),
            current_user: dict[str, str] = Depends(get_current_user),
        ) -> ApiResponse[dict[str, int]]:
            data = StatisticsService(db).get_project_overview(project_code)
            return ApiResponse(data=data)
        """
    ),
    "app/utils/pagination.py": code_block(
        """
        from math import ceil


        def build_page_payload(page: int, page_size: int, total: int) -> dict[str, int]:
            page_count = ceil(total / page_size) if page_size > 0 else 0
            return {
                "page": page,
                "page_size": page_size,
                "total": total,
                "page_count": page_count,
            }


        def clamp_page(page: int) -> int:
            if page <= 1:
                return 1
            return page


        def clamp_page_size(page_size: int) -> int:
            if page_size <= 0:
                return 20
            if page_size > 100:
                return 100
            return page_size
        """
    ),
    "app/utils/date_range.py": code_block(
        """
        from datetime import datetime, timedelta


        def resolve_month_range(reference: datetime) -> tuple[datetime, datetime]:
            start = reference.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            if start.month == 12:
                next_month = start.replace(year=start.year + 1, month=1)
            else:
                next_month = start.replace(month=start.month + 1)
            end = next_month - timedelta(seconds=1)
            return start, end


        def resolve_week_range(reference: datetime) -> tuple[datetime, datetime]:
            start = reference - timedelta(days=reference.weekday())
            start = start.replace(hour=0, minute=0, second=0, microsecond=0)
            end = start + timedelta(days=6, hours=23, minutes=59, seconds=59)
            return start, end


        def normalize_datetime(value: str | None) -> datetime | None:
            if value is None or not value.strip():
                return None
            return datetime.fromisoformat(value)
        """
    ),
    "app/utils/workflow_seed.py": code_block(
        """
        from app.models.project import WorkflowStageRecord, WorkflowTemplateRecord


        def build_default_new_order_template() -> WorkflowTemplateRecord:
            template = WorkflowTemplateRecord(
                template_name="全新订单标准流程",
                order_type="new",
                is_default=True,
                status="enabled",
            )
            template.stages.extend(
                [
                    WorkflowStageRecord(
                        stage_name="教研任务",
                        role_code="researcher",
                        sort_value=1,
                        requires_file_upload=True,
                    ),
                    WorkflowStageRecord(
                        stage_name="风格稿任务",
                        role_code="coordinator",
                        sort_value=2,
                        can_assign=True,
                        requires_file_upload=True,
                    ),
                    WorkflowStageRecord(
                        stage_name="内页任务",
                        role_code="coordinator",
                        sort_value=3,
                        can_assign=True,
                        requires_file_upload=True,
                        requires_validation=True,
                    ),
                    WorkflowStageRecord(
                        stage_name="归档任务",
                        role_code="coordinator",
                        sort_value=4,
                        triggers_package=True,
                    ),
                ]
            )
            return template


        def build_default_iteration_template() -> WorkflowTemplateRecord:
            template = WorkflowTemplateRecord(
                template_name="迭代订单标准流程",
                order_type="iteration",
                is_default=False,
                status="enabled",
            )
            template.stages.extend(
                [
                    WorkflowStageRecord(
                        stage_name="迭代信息确认",
                        role_code="planner",
                        sort_value=1,
                    ),
                    WorkflowStageRecord(
                        stage_name="设计资料重整",
                        role_code="coordinator",
                        sort_value=2,
                        can_assign=True,
                        requires_file_upload=True,
                    ),
                    WorkflowStageRecord(
                        stage_name="成品归档",
                        role_code="coordinator",
                        sort_value=3,
                        triggers_package=True,
                    ),
                ]
            )
            return template
        """
    ),
    "app/seed/bootstrap.py": code_block(
        """
        from sqlalchemy.orm import Session

        from app.core.security import get_password_hash
        from app.models.project import ProjectRecord
        from app.models.user import UserAccount
        from app.repositories.project_repository import ProjectRepository
        from app.repositories.user_repository import UserRepository
        from app.utils.workflow_seed import (
            build_default_iteration_template,
            build_default_new_order_template,
        )


        def bootstrap_system(db: Session) -> dict[str, int]:
            user_repository = UserRepository(db)
            project_repository = ProjectRepository(db)
            created_users = 0
            created_projects = 0

            if user_repository.get_by_email("admin@example.com") is None:
                user_repository.add(
                    UserAccount(
                        email="admin@example.com",
                        password_hash=get_password_hash("123456"),
                        name="系统管理员",
                        role_code="admin",
                        enabled=True,
                    )
                )
                created_users += 1

            if project_repository.get_by_code("DESIGN-DELIVERY") is None:
                project = ProjectRecord(
                    project_code="DESIGN-DELIVERY",
                    project_name="设计交付演示项目",
                    description="默认演示项目",
                    status="enabled",
                )
                project.workflow_templates.append(build_default_new_order_template())
                project.workflow_templates.append(build_default_iteration_template())
                project_repository.add(project)
                created_projects += 1

            project_repository.save()
            return {
                "created_users": created_users,
                "created_projects": created_projects,
            }
        """
    ),
    "app/tasks/packaging.py": code_block(
        """
        from datetime import datetime
        from uuid import uuid4

        from sqlalchemy.orm import Session

        from app.models.course import CourseAttachment, CourseLog
        from app.repositories.course_repository import CourseRepository


        class PackagingTaskRunner:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.repository = CourseRepository(db)

            def run_once(self, task_code: str) -> dict[str, str]:
                task = self.repository.get_by_task_code(task_code)
                if task is None:
                    raise ValueError("course_not_found")
                archive_name = f"{task.task_code}_{task.version_no}_{uuid4().hex[:8]}.zip"
                task.attachments.append(
                    CourseAttachment(
                        file_name=archive_name,
                        category="archive",
                        uploaded_by="system",
                        uploaded_at=datetime.utcnow(),
                    )
                )
                task.logs.append(
                    CourseLog(
                        actor_name="system",
                        action_name="自动打包完成",
                        detail_text=f"生成归档文件 {archive_name}",
                        created_at=datetime.utcnow(),
                    )
                )
                task.status = "archived"
                task.quality_check = "打包成功"
                task.archived_at = datetime.utcnow()
                self.repository.save()
                return {
                    "task_code": task.task_code,
                    "archive_name": archive_name,
                    "status": task.status,
                }
        """
    ),
    "app/tasks/reminder.py": code_block(
        """
        from datetime import datetime, timedelta

        from sqlalchemy.orm import Session

        from app.models.course import CourseLog
        from app.repositories.course_repository import CourseRepository


        class ReminderTaskRunner:
            def __init__(self, db: Session) -> None:
                self.db = db
                self.repository = CourseRepository(db)

            def mark_overdue_tasks(self) -> int:
                tasks, _total = self.repository.list(limit=5000)
                count = 0
                today = datetime.utcnow().date()
                for task in tasks:
                    if task.overall_due_date is None:
                        continue
                    if task.status == "archived":
                        continue
                    if task.overall_due_date < today and not task.overdue:
                        task.overdue = True
                        task.logs.append(
                            CourseLog(
                                actor_name="system",
                                action_name="逾期提醒",
                                detail_text="任务已超过总交付日期",
                                created_at=datetime.utcnow(),
                            )
                        )
                        count += 1
                self.repository.save()
                return count

            def mark_due_soon_tasks(self) -> int:
                tasks, _total = self.repository.list(limit=5000)
                count = 0
                threshold = datetime.utcnow().date() + timedelta(days=2)
                for task in tasks:
                    if task.overall_due_date is None:
                        continue
                    if task.status == "archived":
                        continue
                    if task.overall_due_date <= threshold:
                        task.logs.append(
                            CourseLog(
                                actor_name="system",
                                action_name="临近交付提醒",
                                detail_text="任务将在两天内到期",
                                created_at=datetime.utcnow(),
                            )
                        )
                        count += 1
                self.repository.save()
                return count
        """
    ),
    "app/tests/test_workflow_service.py": code_block(
        """
        from datetime import date

        from app.models.course import CourseStage, CourseTask
        from app.schemas.workflow import AdvanceWorkflowRequest
        from app.services.workflow_service import WorkflowService


        def build_task() -> CourseTask:
            task = CourseTask(
                task_code="KC-2026-001",
                title="测试课件",
                series="新业务演示",
                subject="语文",
                education_stage="高中",
                grade="高一",
                volume="必修上",
                textbook="统编版",
                order_type="全新订单",
                version_no="v1.0",
                status="research",
                current_owner="教研老师 · 赵老师",
                coordinator="林薇",
                research_owner="赵老师",
                research_due_date=date(2026, 5, 30),
                final_due_date=date(2026, 6, 6),
                overall_due_date=date(2026, 6, 8),
            )
            task.stages.extend(
                [
                    CourseStage(stage_key="research", label="教研任务", owner_role="researcher", owner_name="赵老师", deliverable="资料", state="active"),
                    CourseStage(stage_key="style", label="风格稿任务", owner_role="coordinator", owner_name="林薇", deliverable="风格稿", state="pending"),
                    CourseStage(stage_key="page", label="内页任务", owner_role="coordinator", owner_name="林薇", deliverable="内页", state="pending"),
                    CourseStage(stage_key="archive", label="归档任务", owner_role="coordinator", owner_name="林薇", deliverable="压缩包", state="pending"),
                ]
            )
            return task


        def test_advance_research_to_style_dispatch() -> None:
            task = build_task()
            updated = WorkflowService().advance(task, "赵老师", "提交教研资料")
            assert updated.status == "pendingStyleDispatch"
            assert updated.current_owner == "设计统筹 · 林薇"


        def test_advance_to_archive_generates_package() -> None:
            task = build_task()
            service = WorkflowService()
            for actor in ["赵老师", "林薇", "唐婧", "林薇", "江栩", "林薇", "system"]:
                task = service.advance(task, actor)
            assert task.status == "archived"
            assert task.quality_check == "打包成功"
            assert len(task.attachments) >= 1
        """
    ),
    "app/tests/test_project_service.py": code_block(
        """
        from app.schemas.project import FieldConfigPayload, WorkflowStagePayload, WorkflowTemplatePayload
        from app.services.validation_service import ValidationService


        def test_validate_field_configs() -> None:
            payloads = [
                FieldConfigPayload(
                    field_key="series",
                    field_name="系列",
                    field_type="text",
                    required=True,
                    searchable=True,
                    sort_value=1,
                    status="enabled",
                ),
                FieldConfigPayload(
                    field_key="subject",
                    field_name="学科",
                    field_type="select",
                    required=True,
                    searchable=True,
                    sort_value=2,
                    status="enabled",
                ),
            ]
            ValidationService().validate_field_configs(payloads)


        def test_validate_workflow_templates() -> None:
            payloads = [
                WorkflowTemplatePayload(
                    template_name="标准流程",
                    order_type="new",
                    is_default=True,
                    status="enabled",
                    stages=[
                        WorkflowStagePayload(stage_name="教研任务", role_code="researcher", sort_value=1),
                        WorkflowStagePayload(stage_name="风格稿任务", role_code="coordinator", sort_value=2),
                    ],
                )
            ]
            ValidationService().validate_workflow_templates(payloads)
        """
    ),
}


def read_existing_files(paths: Iterable[Path]) -> dict[str, str]:
    result: dict[str, str] = {}
    for path in paths:
        relative = path.relative_to(OUTPUT_ROOT / "backend_source_30_pages").as_posix()
        result[relative] = path.read_text(encoding="utf-8")
    return result


def build_all_modules() -> dict[str, str]:
    modules = read_existing_files(EXISTING_FILES)
    modules.update(NEW_MODULES)
    return dict(sorted(modules.items(), key=lambda item: item[0]))


def write_source_tree(modules: dict[str, str]) -> None:
    SOURCE_ROOT.mkdir(parents=True, exist_ok=True)
    for relative_path, content in modules.items():
        path = SOURCE_ROOT / relative_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")


def flatten_lines(modules: dict[str, str]) -> list[str]:
    lines: list[str] = []
    for relative_path, content in modules.items():
        lines.append(f"File: {relative_path}")
        lines.extend(content.splitlines())
        lines.append("")
    while lines and lines[-1] == "":
        lines.pop()
    return lines


def split_evenly(lines: list[str], page_count: int) -> list[list[str]]:
    if len(lines) < page_count:
        raise ValueError("代码行数不足以拆分为目标页数")
    base = len(lines) // page_count
    remainder = len(lines) % page_count
    chunks: list[list[str]] = []
    start = 0
    for index in range(page_count):
        size = base + (1 if index < remainder else 0)
        end = start + size
        chunks.append(lines[start:end])
        start = end
    return chunks


def build_docx(page_chunks: list[list[str]], output_path: Path) -> None:
    doc = Document(REFERENCE_DOC)
    body = doc._body._body
    sect_pr = body[-1]
    for child in list(body)[:-1]:
        body.remove(child)
    for page_index, chunk in enumerate(page_chunks):
        paragraph = doc.add_paragraph(style="HTML Preformatted")
        format_rule = paragraph.paragraph_format
        format_rule.space_before = Pt(0)
        format_rule.space_after = Pt(0)
        format_rule.left_indent = Pt(0)
        format_rule.first_line_indent = Pt(0)
        run = paragraph.add_run()
        run.font.name = "等线"
        run._element.rPr.rFonts.set(qn("w:ascii"), "等线")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "等线")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run._element.rPr.rFonts.set(qn("w:cs"), "等线")
        run.font.size = Pt(9)
        for line_index, line in enumerate(chunk):
            run.add_text(line if line else " ")
            if line_index < len(chunk) - 1:
                run.add_break(WD_BREAK.LINE)
        if page_index < len(page_chunks) - 1:
            run.add_break(WD_BREAK.PAGE)
    doc.save(output_path)


def main() -> None:
    modules = build_all_modules()
    write_source_tree(modules)
    lines = flatten_lines(modules)
    page_chunks = split_evenly(lines, page_count=60)
    output_60 = OUTPUT_ROOT / "设计交付系统_后端源代码60页.docx"
    output_main = OUTPUT_ROOT / "设计交付系统_源代码60页_紧凑版.docx"
    build_docx(page_chunks, output_60)
    build_docx(page_chunks, output_main)
    print(f"modules={len(modules)}")
    print(f"lines={len(lines)}")
    print(f"page1_lines={len(page_chunks[0])}")
    print(f"page60_lines={len(page_chunks[-1])}")
    print(output_60)


if __name__ == "__main__":
    main()
